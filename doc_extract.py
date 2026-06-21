"""
doc_extract.py
==========================================
Utility to extract text from various document formats:
PDF, DOCX, and images (scans) via OCR.

If the OCR library (pytesseract) or tesseract binary is not available,
image extraction returns a warning message instead of crashing.
==========================================
"""
import os

import pdfplumber
from docx import Document as DocxDocument

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "png", "jpg", "jpeg"}


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF (text-based PDF, not scanned image)."""
    text_parts = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
    except Exception as e:
        return f"[Failed to read PDF: {e}]"

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        return (
            "[This PDF appears to be a scanned image without readable text. "
            "Upload it as an image (PNG/JPG) for OCR, or use a PDF exported directly from Word.]"
        )
    return full_text


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file, including table content."""
    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        return f"[Failed to read DOCX: {e}]"

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text_from_image(filepath: str) -> str:
    """Extract text from an image (document scan) using OCR."""
    if not OCR_AVAILABLE:
        return (
            "[OCR is not available on this server. Install with: "
            "`pip install pytesseract Pillow` and make sure the `tesseract` binary "
            "is installed on the system.]"
        )
    try:
        image = Image.open(filepath)
        # lang='ind+eng' to support Indonesian and English OCR
        text = pytesseract.image_to_string(image, lang="ind+eng")
        return text.strip() or "[No readable text was found in this image.]"
    except Exception as e:
        return f"[Failed to read image with OCR: {e}]"


def extract_text(filepath: str, filename: str) -> str:
    """Main dispatcher: select extraction method based on file extension."""
    ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""

    if ext == "pdf":
        return extract_text_from_pdf(filepath)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(filepath)
    elif ext in ("png", "jpg", "jpeg"):
        return extract_text_from_image(filepath)
    else:
        return f"[File format .{ext} is not supported.]"
