"""
doc_export.py
==========================================
Convert teks draft (plain text, paragraf dipisah baris kosong) menjadi
file .docx atau .pdf yang rapi, siap di-download user.
==========================================
"""
import io

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY


def _split_paragraphs(text: str) -> list:
    """Pisahkan teks jadi paragraf berdasarkan baris kosong."""
    raw_parts = [p.strip() for p in text.split("\n\n")]
    return [p for p in raw_parts if p]


def text_to_docx_bytes(title: str, text: str) -> bytes:
    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    heading = doc.add_heading(title, level=1)

    for paragraph in _split_paragraphs(text):
        # Baris dalam satu paragraf tetap dipisah (misal alamat bertingkat)
        lines = paragraph.split("\n")
        p = doc.add_paragraph(lines[0])
        for line in lines[1:]:
            p.add_run("\n" + line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def text_to_pdf_bytes(title: str, text: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle", parent=styles["Heading1"], fontSize=16, spaceAfter=18,
    )
    body_style = ParagraphStyle(
        "BodyStyle", parent=styles["Normal"], fontSize=11, leading=16,
        alignment=TA_JUSTIFY, spaceAfter=12,
    )

    elements = [Paragraph(title, title_style), Spacer(1, 6)]

    for paragraph in _split_paragraphs(text):
        # reportlab Paragraph butuh <br/> untuk line break, bukan \n
        html_safe = (
            paragraph.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )
        html_safe = html_safe.replace("\n", "<br/>")
        elements.append(Paragraph(html_safe, body_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.read()
